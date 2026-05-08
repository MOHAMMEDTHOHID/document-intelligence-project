import os
import PyPDF2
from docx import Document as DocxDocument
from pathlib import Path

class DocumentProcessor:
    """Handle document processing and text extraction"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text(self, file_path):
        """Extract text from various document formats"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf(file_path)
        elif file_ext == '.docx':
            return self._extract_docx(file_path)
        elif file_ext == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_pdf(self, file_path):
        """Extract text from PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                    text += f"\n--- Page {page_num + 1} ---\n"
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
        
        return text.strip()
    
    def _extract_docx(self, file_path):
        """Extract text from DOCX"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
        
        return text.strip()
    
    def _extract_txt(self, file_path):
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            raise Exception(f"Error extracting TXT: {str(e)}")
        
        return text.strip()
    
    def chunk_text(self, text, chunk_size=1000, overlap=100):
        """Split text into chunks for processing"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap
        
        return chunks
    
    def clean_text(self, text):
        """Clean and normalize text"""
        # Remove extra whitespace
        text = ' '.join(text.split())
        # Remove special characters (keep basic punctuation)
        text = ''.join(c for c in text if c.isalnum() or c.isspace() or c in '.,!?;:')
        return text
